from __future__ import annotations

from datetime import datetime
import os
import os.path
from pathlib import Path
from tempfile import mkdtemp
from typing import Dict

from makevideopreviews.errors import DependencyError
from makevideopreviews.models import AppConfig, ExtractionResult, FolderJob, PreviewScope, VideoProbe
from makevideopreviews.utils import atomic_docx_path, format_resolution, safe_output_name, seconds_to_hms


def ensure_generate_dependencies():
    try:
        from PIL import Image, ImageDraw, ImageFont, ImageOps
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ModuleNotFoundError as exc:
        raise DependencyError(
            "Missing Python dependency for document generation: %s. Install with `python3 -m pip install -e .`."
            % exc.name
        )
    return Document, Image, ImageDraw, ImageFont, ImageOps, WD_ORIENT, WD_ALIGN_PARAGRAPH, Inches, Pt


def write_docx(job: FolderJob, results: Dict[Path, ExtractionResult], config: AppConfig) -> Path:
    Document, Image, ImageDraw, ImageFont, ImageOps, WD_ORIENT, WD_ALIGN_PARAGRAPH, Inches, Pt = ensure_generate_dependencies()

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Helvetica"
    normal_style.font.size = Pt(10)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)

    _add_cover_page(document, job, config, WD_ALIGN_PARAGRAPH, Pt)

    normalized_root = Path(mkdtemp(prefix="mvp_docx_"))
    try:
        filename_prefixes = _common_filename_prefixes(job.videos)
        wrote_content = False
        for index, probe in enumerate(job.probes):
            result = results.get(probe.path)
            if result is None or result.error:
                if wrote_content:
                    document.add_page_break()
                else:
                    document.add_page_break()
                    wrote_content = True
                _add_video_section_header(document, probe, Pt)
                if probe.problem:
                    paragraph = document.add_paragraph()
                    paragraph.add_run("Metadata warning: %s" % probe.problem).italic = True
                paragraph = document.add_paragraph()
                paragraph.add_run("Could not extract preview frames.").bold = True
                if result and result.error:
                    document.add_paragraph(result.error)
                continue

            sheets = _build_contact_sheet_pages(
                Image,
                ImageDraw,
                ImageFont,
                ImageOps,
                job,
                probe,
                result,
                config,
                normalized_root,
                filename_prefixes,
            )
            for sheet_index, sheet_path in enumerate(sheets):
                if wrote_content:
                    document.add_page_break()
                else:
                    document.add_page_break()
                    wrote_content = True
                normalized = _normalize_for_docx(Image, sheet_path, normalized_root / ("%s_docx_%s" % (probe.path.stem, sheet_path.name)))
                available_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
                available_height = section.page_height.inches - section.top_margin.inches - section.bottom_margin.inches
                picture_width = _fit_picture_width_inches(normalized, available_width, available_height)
                document.add_picture(str(normalized), width=Inches(min(config.page_width_inch, picture_width)))
                picture_paragraph = document.paragraphs[-1]
                picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                picture_paragraph.paragraph_format.space_before = Pt(0)
                picture_paragraph.paragraph_format.space_after = Pt(0)

        tmp_path = atomic_docx_path(job.output_path)
        document.save(str(tmp_path))
        os.replace(str(tmp_path), str(job.output_path))
        return job.output_path
    finally:
        for child in normalized_root.glob("*"):
            child.unlink(missing_ok=True)
        normalized_root.rmdir()


def _add_cover_page(document, job, config, WD_ALIGN_PARAGRAPH, Pt):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Preview: %s" % job.folder.name)
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.add_run("Generated: %s" % datetime.now().strftime("%Y-%m-%d %H:%M:%S")).italic = True

    details = [
        "Folder: %s" % job.folder,
        "Videos: %d" % len(job.videos),
        "Estimated output size: %s" % _humanize_estimate(job.estimated_bytes),
        "Interval: %ss" % config.interval,
        "Max width: %spx" % config.max_width,
        "JPEG quality (-q:v): %s" % config.quality,
        "Page width: %.1f in" % config.page_width_inch,
    ]
    for detail in details:
        document.add_paragraph(detail)

    if job.warnings:
        warning_header = document.add_paragraph()
        warning_run = warning_header.add_run("Warnings")
        warning_run.bold = True
        for warning in job.warnings:
            document.add_paragraph(warning)


def _add_video_section_header(document, probe: VideoProbe, Pt):
    heading = document.add_paragraph()
    heading_run = heading.add_run(probe.path.name)
    heading_run.bold = True
    heading_run.font.size = Pt(14)

    metadata = [
        "Duration: %s" % (seconds_to_hms(probe.duration) if probe.duration else "unknown"),
        "Resolution: %s" % format_resolution(probe.width, probe.height),
        "Frame rate: %s fps" % ("%.2f" % probe.fps if probe.fps else "unknown"),
        "Expected frames: %s" % probe.estimated_frames,
    ]
    for detail in metadata:
        paragraph = document.add_paragraph()
        paragraph.add_run(detail)


def _relative_folder_label(root: Path, folder: Path) -> str:
    try:
        relative = folder.relative_to(root)
        return str(relative) if str(relative) != "." else root.name
    except ValueError:
        return str(folder)


def _build_contact_sheet_pages(Image, ImageDraw, ImageFont, ImageOps, job, probe, result, config, output_dir: Path, filename_prefixes):
    page_width_px = int(config.page_width_inch * 220)
    page_height_px = int(page_width_px / (10.3 / 7.8))
    title_height = max(70, int(page_height_px * 0.07))
    inner_margin = max(20, int(page_width_px * 0.018))
    gap = max(10, int(page_width_px * 0.006))
    thumb_aspect = _thumbnail_aspect(probe)
    cols, rows, thumb_w, thumb_h = _choose_grid(page_width_px, page_height_px, title_height, inner_margin, gap, thumb_aspect)
    per_page = cols * rows
    font = _load_font(ImageFont, max(14, int(thumb_h * 0.09)))
    title_font = _load_font(ImageFont, max(18, int(title_height * 0.38)))

    relative_folder = _relative_folder_label(job.folder, probe.path.parent)
    overlay_name = _overlay_name_for_probe(probe, filename_prefixes)
    pages = []
    frames = list(zip(result.frame_paths, result.timestamps))
    total_pages = max(1, (len(frames) + per_page - 1) // per_page)
    for page_index in range(total_pages):
        chunk = frames[page_index * per_page:(page_index + 1) * per_page]
        canvas = Image.new("RGB", (page_width_px, page_height_px), color=(248, 246, 240))
        draw = ImageDraw.Draw(canvas, "RGBA")
        _draw_page_header(draw, page_width_px, title_height, title_font, relative_folder, probe.path.name, page_index + 1, total_pages)

        for thumb_index, (frame_path, timestamp) in enumerate(chunk):
            row = thumb_index // cols
            col = thumb_index % cols
            x = inner_margin + col * (thumb_w + gap)
            y = title_height + inner_margin + row * (thumb_h + gap)
            with Image.open(str(frame_path)) as raw_image:
                thumb = ImageOps.fit(raw_image.convert("RGB"), (thumb_w, thumb_h), method=Image.Resampling.LANCZOS)
            canvas.paste(thumb, (x, y))
            _draw_thumb_overlay(draw, font, x, y, thumb_w, thumb_h, overlay_name, seconds_to_hms(timestamp))

        out_path = output_dir / ("%s_sheet_%03d.jpg" % (safe_output_name(probe.path.name), page_index))
        canvas.save(str(out_path), format="JPEG", quality=92, optimize=True, progressive=False)
        pages.append(out_path)

    return pages


def _choose_grid(page_width_px, page_height_px, title_height, inner_margin, gap, thumb_aspect):
    best = None
    usable_width = page_width_px - (2 * inner_margin)
    usable_height = page_height_px - title_height - (2 * inner_margin)
    for cols in range(6, 9):
        for rows in range(5, 9):
            max_width = int((usable_width - gap * (cols - 1)) / cols)
            max_height = int((usable_height - gap * (rows - 1)) / rows)
            thumb_width = min(max_width, int(max_height * thumb_aspect))
            thumb_height = int(thumb_width / thumb_aspect)
            if thumb_width <= 0 or thumb_height <= 0:
                continue
            coverage = thumb_width * thumb_height * cols * rows
            score = (coverage, -abs(cols - 7), -abs(rows - 6))
            if best is None or score > best[0]:
                best = (score, cols, rows, thumb_width, thumb_height)
    if best is None:
        return 6, 5, 240, 135
    return best[1], best[2], best[3], best[4]


def _draw_page_header(draw, page_width_px, title_height, font, folder_label, video_name, page_number, total_pages):
    draw.rectangle([0, 0, page_width_px, title_height], fill=(34, 42, 53, 255))
    title = "%s | %s | page %s/%s" % (folder_label, video_name, page_number, total_pages)
    draw.text((26, max(12, int(title_height * 0.24))), title, font=font, fill=(255, 255, 255, 255))


def _draw_thumb_overlay(draw, font, x, y, thumb_w, thumb_h, file_name, timecode):
    overlay_height = max(22, int(thumb_h * 0.17))
    draw.rectangle([x, y, x + thumb_w, y + overlay_height], fill=(0, 0, 0, 175))
    inner_padding = 8
    timecode_width = font.getlength(timecode)
    time_x = x + thumb_w - inner_padding - timecode_width
    text_y = y + max(4, int((overlay_height - font.size) / 2))
    draw.text((time_x, text_y), timecode, font=font, fill=(255, 255, 255, 255))
    available_name_width = max(0, int(time_x - (x + inner_padding + 10)))
    name = _clip_text(file_name, available_name_width, font)
    draw.text((x + inner_padding, text_y), name, font=font, fill=(255, 255, 255, 255))


def _clip_text(text, max_width, font):
    if font.getlength(text) <= max_width:
        return text
    clipped = text
    while clipped and font.getlength(clipped + "...") > max_width:
        clipped = clipped[:-1]
    return clipped + "..." if clipped else "..."


def _thumbnail_aspect(probe: VideoProbe) -> float:
    if probe.width and probe.height and probe.height > 0:
        return float(probe.width) / float(probe.height)
    return 16.0 / 9.0


def _load_font(ImageFont, size):
    for font_name in ("Helvetica.ttc", "Arial Unicode.ttf", "Arial.ttf", "DejaVuSans.ttf"):
        try:
            return ImageFont.truetype(font_name, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def _common_filename_prefixes(video_paths):
    grouped = {}
    for video_path in video_paths:
        grouped.setdefault(video_path.parent, []).append(video_path.stem)

    prefixes = {}
    for folder, stems in grouped.items():
        prefix = os.path.commonprefix(stems)
        prefix = prefix.rstrip(" _-.")
        prefixes[folder] = prefix if len(prefix) >= 4 else ""
    return prefixes


def _overlay_name_for_probe(probe: VideoProbe, prefixes):
    stem = probe.path.stem
    prefix = prefixes.get(probe.path.parent) or ""
    shortened = stem[len(prefix):] if prefix and stem.startswith(prefix) else stem
    shortened = shortened.lstrip(" _-.")
    if not shortened:
        shortened = stem
    return "%s%s" % (shortened, probe.path.suffix.lower())


def _fit_picture_width_inches(image_path: Path, max_width_inch: float, max_height_inch: float) -> float:
    from PIL import Image

    height_budget = max(1.0, max_height_inch - 0.35)
    width_budget = max(1.0, max_width_inch - 0.15)

    with Image.open(str(image_path)) as image:
        width, height = image.size
    if width <= 0 or height <= 0:
        return width_budget
    aspect = float(width) / float(height)
    height_limited_width = height_budget * aspect
    return min(width_budget, height_limited_width)


def _normalize_for_docx(Image, src: Path, dest: Path) -> Path:
    with Image.open(str(src)) as image:
        normalized = image.convert("RGB")
        normalized.save(str(dest), format="JPEG", quality=92, optimize=True, progressive=False)
    return dest


def _humanize_estimate(size: int) -> str:
    if size <= 0:
        return "unknown"
    if size < 1024:
        return "%d B" % size
    if size < 1024 * 1024:
        return "%.1f KB" % (size / 1024.0)
    if size < 1024 * 1024 * 1024:
        return "%.1f MB" % (size / (1024.0 * 1024.0))
    return "%.1f GB" % (size / (1024.0 * 1024.0 * 1024.0))
